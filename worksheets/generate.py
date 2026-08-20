# -*- coding: utf-8 -*-
"""
문해력 15분 — 손으로 쓰는 학습지(docx) 생성기

한 주제 = 1주(Day1~5)로 구성한다 (Day1~2 학습, Day3~4 확인, Day5 종합복습).
웹앱(data/content/wXX.js)과 문항·정답 기준이 동일하다. 2단 배치와 인라인 빈칸으로
하루 분량을 15분 안에 끝낼 수 있게 압축했다.

사용법:
    python generate.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "맑은 고딕"
ACCENT = RGBColor(0x3D, 0x4B, 0xC7)
GRAY = RGBColor(0x6B, 0x6F, 0x85)
LIGHTGRAY = "F2F3F8"
BLANK = "＿＿＿＿＿＿"

# =========================================================================
# 데이터 (data/content/w01.js 의 "w01d01" / "w01d02" 내용과 동일)
# =========================================================================

WEEK1 = {
    "week": 1, "unit": "Ⅰ. 생존 문해력", "topic": "공문서·행정 문서 읽기",

    "d1": {
        "tag": "학습①", "intro": "6개 어휘부터 뜻을 읽고, 예문의 빈칸에 낱말을 직접 써서 문장을 완성해 보세요.",
        "words": [
            {"term": "신청", "hanja": "申請·펼신·청할청", "definition": "단체·기관에 어떤 일을 알려 청구함.", "example": "그는 도서관 열람실 이용을 {}했다."},
            {"term": "발급", "hanja": "發給·필발·줄급", "definition": "증명서 따위를 발행하여 내줌.", "example": "주민센터는 증명서를 즉시 {}해 준다."},
            {"term": "제출", "hanja": "提出·끌제·날출", "definition": "문안·의견·서류 따위를 내놓음.", "example": "결석계는 담임 선생님께 직접 {}한다."},
            {"term": "첨부", "hanja": "添附·더할첨·붙을부", "definition": "문서나 물건 따위를 더 보탬.", "example": "신청서에는 신분증 사본을 {}한다."},
            {"term": "기한", "hanja": "期限·기약할기·한할한", "definition": "미리 한정하여 놓은 시기.", "example": "서류 제출 {}을 넘기면 취소된다."},
            {"term": "유효", "hanja": "有效·있을유·효험효", "definition": "보람이나 효과가 있음. (↔무효)", "example": "이 쿠폰은 이달 말까지만 {}하다."},
        ],
        "checkpoints": [
            {"options": ("신청", "발급"), "template": "여권용 사진을 들고 가서 여권 {0}을 했다. 2주 뒤 여권이 {1}되었다는 문자를 받았다."},
            {"options": ("제출", "첨부"), "template": "봉사활동 확인서를 {0}할 때에는 활동 사진을 함께 {1}해야 인정받는다."},
            {"options": ("기한", "유효"), "template": "이 신분증은 12월까지만 {1}하며, 그 {0}이 지나면 다시 발급받아야 한다."},
        ],
        "quote": {"text": "천릿길도 한 걸음부터 시작된다.", "author": "한국 속담"},
    },

    "d2": {
        "tag": "학습②", "intro": "어제 배운 6개에 이어, 나머지 6개 어휘를 익혀 봅시다.",
        "words": [
            {"term": "열람", "hanja": "閱覽·볼열·볼람", "definition": "문서 따위를 죽 훑어보거나 조사함.", "example": "이 서류는 허락을 받아야 {}할 수 있다."},
            {"term": "통보", "hanja": "通報·통할통·알릴보", "definition": "통지하여 보고함.", "example": "합격자에게는 문자로 결과를 {}한다."},
            {"term": "반려", "hanja": "返戾·돌이킬반·어그러질려", "definition": "제출한 문서를 되돌려 보냄.", "example": "서명이 빠진 신청서는 결국 {}되었다."},
            {"term": "사본", "hanja": "寫本·베낄사·근본본", "definition": "원본을 그대로 옮긴 문서. (↔원본)", "example": "신분증 {} 한 부를 제출해 주세요."},
            {"term": "갱신", "hanja": "更新·다시갱·새신", "definition": "이미 있던 것을 고쳐 새롭게 함.", "example": "운전면허증은 주기적으로 {}한다."},
            {"term": "공고", "hanja": "公告·공평할공·알릴고", "definition": "기관이 어떤 사항을 세상에 널리 알림.", "example": "채용 {}에는 지원 자격이 나와 있다."},
        ],
        "checkpoints": [
            {"options": ("열람", "통보"), "template": "개인정보 서류는 담당자만 {0}할 수 있고, 결과는 문자로 {1}된다."},
            {"options": ("반려", "사본"), "template": "서명이 빠진 서류는 {0}되었다. 재접수할 때는 신분증 {1}을 함께 냈다."},
            {"options": ("갱신", "공고"), "template": "운전면허증은 정기적으로 {0}해야 하며, 채용 {1}에는 지원 자격이 나와 있다."},
        ],
        "confusable": {
            "left": {"term": "결제", "hanja": "決濟", "definition": "돈을 주고받아 매매 거래를 끝맺음.", "example": "카드로 결제했다."},
            "right": {"term": "결재", "hanja": "決裁", "definition": "책임자가 안건을 검토하여 허가함.", "example": "부장님의 결재를 받았다."},
            "tip": "값을 '치르면' 결제, 서류를 '검토받으면' 결재!",
        },
        "quote": {"text": "우리는 반복적으로 행하는 대로 된다. 탁월함은 행동이 아니라 습관이다.", "author": "아리스토텔레스"},
    },

    "d3": {
        "tag": "확인①", "intro": "이틀 동안 배운 12개 어휘를 세 가지 방식으로 확인해 봅시다.",
        "sections": [
            {"title": "① 초성 힌트 — 빈칸에 낱말을 쓰세요", "type": "fill", "items": [
                {"prompt": "[ㅅㅊ] 알려 청구함", "hint": "동아리 가입 ___을 오늘까지 받는다."},
                {"prompt": "[ㅂㄱ] 발행하여 내줌", "hint": "무인 발급기에서도 ___받을 수 있다."},
                {"prompt": "[ㅊㅂ] 더 보탬", "hint": "신청서에 사진을 ___해서 냈다."},
                {"prompt": "[ㅂㄹ] 되돌려 보냄", "hint": "서명이 빠진 서류는 ___되었다."},
            ]},
            {"title": "② 글자 카드 조합 — 기·한·유·효·통·보·갱·신", "type": "fill", "items": [
                {"prompt": "미리 한정하여 놓은 시기"},
                {"prompt": "보람이나 효과가 있음"},
                {"prompt": "통지하여 보고함"},
                {"prompt": "고쳐 새롭게 함"},
            ]},
            {"title": "③ 문장에 어울리는 어휘 고르기", "type": "choice", "items": [
                {"prompt": "사진과 신분증을 가지고 가서 (　) 한다.", "options": ["신청", "제출"]},
                {"prompt": "확인서는 담당 선생님께 (　) 한다.", "options": ["첨부", "제출"]},
                {"prompt": "관계자 외에는 (　) 할 수 없다.", "options": ["열람", "통보"]},
                {"prompt": "채용 (　)에는 지원 자격이 나와 있다.", "options": ["공고", "사본"]},
            ]},
        ],
        "quote": {"text": "아는 것을 안다고 하고, 모르는 것을 모른다고 하는 것 — 그것이 진짜 아는 것이다.", "author": "공자"},
    },

    "d4": {
        "tag": "확인②", "intro": "밑줄 친 어휘의 뜻을 확인하고, 실제 안내문 속에서 어떻게 쓰이는지 살펴봅시다.",
        "sections": [
            {"title": "④ 밑줄 뜻으로 알맞은 것 고르기", "type": "choice4", "items": [
                {"prompt": "예산안을 학생부에 제출했다.", "options": ["서류를 냄", "허가함", "훑어봄", "대신함"]},
                {"prompt": "자격증은 유효 기간이 5년이다.", "options": ["효력 있음", "거짓 판명", "기간 끝남", "새로 만듦"]},
                {"prompt": "합격 여부는 문자로 통보한다.", "options": ["직접 확인", "통지·보고", "서류 첨부", "기한 지정"]},
                {"prompt": "원본이 없으면 사본이라도 가져오세요.", "options": ["옮긴 문서", "새 문서", "지난 문서", "빠진 문서"]},
            ]},
        ],
        "passage": {
            "title": "⑤ 안내문을 읽고 답하세요",
            "body": [
                "[교내 방과후학교 프로그램 신청 안내]",
                "1. 대상: 전교생   2. 기한: 8/25(월)~8/29(금) — 넘기면 접수 불가",
                "3. 방법: 신청서 작성 후 담임 선생님께 제출, 학부모 동의서 첨부",
                "   ※ 서명이 빠진 신청서는 반려됨",
                "4. 발표: 9/1(월), 결과는 문자로 통보   5. 서류는 교무실에서만 열람, 이번 학기만 유효",
                "6. 자세한 내용은 교내 공고문 참고",
            ],
            "items": [
                {"prompt": "신청서와 함께 반드시 내야 하는 것은?", "options": ["학부모 동의서", "성적표", "자기소개서", "등본"]},
                {"prompt": "신청 기한을 넘기면?", "options": ["자동 연장", "접수 불가", "문자 안내", "서류 추가"]},
                {"prompt": "'열람'의 뜻으로 가장 알맞은 것은?", "options": ["폐기함", "복사함", "훑어봄", "재제출"]},
                {"prompt": "이 안내문의 목적은?", "options": ["신청 방법 안내", "프로그램 홍보", "참가비 안내", "강사 모집"]},
                {"prompt": "내용과 일치하지 않는 것은?", "options": ["대상은 전교생", "결과는 이메일 통보", "교무실에서만 열람", "서명 없으면 반려"]},
            ],
        },
        "quote": {"text": "좋은 책을 읽는 것은 과거의 가장 훌륭한 사람들과 대화를 나누는 것과 같다.", "author": "르네 데카르트"},
    },

    "d5": {
        "tag": "종합복습", "intro": "이번 주에 배운 12개 어휘를 새로운 문장으로 마지막으로 점검해 봅시다.",
        "sections": [
            {"title": "뜻을 보고 어휘를 쓰세요", "type": "fill", "items": [
                {"prompt": "증명서 따위를 발행하여 내줌"},
                {"prompt": "문서나 물건 따위를 더 보탬"},
                {"prompt": "보람이나 효과가 있음"},
                {"prompt": "제출한 문서를 처리하지 않고 되돌려 보냄"},
                {"prompt": "원본을 그대로 옮기거나 복사한 문서"},
                {"prompt": "국가·단체가 어떤 사항을 세상에 널리 알림"},
            ]},
            {"title": "문장을 읽고 알맞은 어휘를 고르세요", "type": "choice", "items": [
                {"prompt": "여권용 사진을 들고 가서 여권 (　)을 했다.", "options": ["신청", "제출"]},
                {"prompt": "증명서는 무인 기기에서도 즉시 (　)받을 수 있다.", "options": ["발급", "첨부"]},
                {"prompt": "서류 (　)을 넘기면 접수가 취소된다.", "options": ["기한", "반려"]},
                {"prompt": "이 쿠폰은 이달 말까지만 (　)하다.", "options": ["유효", "사본"]},
                {"prompt": "이 서류는 허락을 받아야 (　)할 수 있다.", "options": ["열람", "갱신"]},
                {"prompt": "합격자에게는 문자로 결과를 (　)한다.", "options": ["통보", "공고"]},
            ]},
        ],
        "quote": {"text": "오늘 할 수 있는 일을 내일로 미루지 마라.", "author": "벤저민 프랭클린"},
    },
}

WEEK2 = {
    "week": 2, "unit": "Ⅰ. 생존 문해력", "topic": "금융·경제 생활",

    "d1": {
        "tag": "학습①", "intro": "용돈을 관리하고 통장을 확인할 때 자주 쓰는 금융 용어예요. 6개 어휘부터 익혀 봅시다.",
        "words": [
            {"term": "예산", "hanja": "豫算·미리예·셈산", "definition": "필요한 비용을 미리 헤아려 계산함. 또는 그 비용.", "example": "이번 달 용돈 {}을 세워 보았다."},
            {"term": "지출", "hanja": "支出·지탱할지·날출", "definition": "어떤 목적을 위하여 돈을 씀. 또는 그 돈.", "example": "이번 달은 교통비 {}이 늘었다."},
            {"term": "저축", "hanja": "貯蓄·쌓을저·쌓을축", "definition": "절약하여 모아 둠.", "example": "그는 매달 용돈의 절반을 {}한다."},
            {"term": "이자", "hanja": "利子·이로울이·아들자", "definition": "돈을 빌려 쓴 대가로 치르는 일정 비율의 돈.", "example": "예금을 하면 {}가 붙는다."},
            {"term": "잔액", "hanja": "殘額·남을잔·액수액", "definition": "쓰고 남은 나머지 돈.", "example": "통장 {}을 확인해 보았다."},
            {"term": "입금", "hanja": "入金·들입·돈금", "definition": "계좌에 돈을 넣음.", "example": "세뱃돈이 통장에 {}되었다."},
        ],
        "checkpoints": [
            {"options": ("예산", "지출"), "template": "이번 달 용돈 {0}을 세웠는데, 생각보다 교통비 {1}이 많았다."},
            {"options": ("저축", "이자"), "template": "매달 용돈의 절반을 {0}했더니, {1}까지 붙어 목돈이 되었다."},
            {"options": ("잔액", "입금"), "template": "세뱃돈이 통장에 {1}되어, {0}을 확인해 보니 크게 늘어 있었다."},
        ],
        "quote": {"text": "티끌 모아 태산.", "author": "한국 속담"},
    },

    "d2": {
        "tag": "학습②", "intro": "어제 배운 6개에 이어, 나머지 6개 어휘를 익혀 봅시다.",
        "words": [
            {"term": "인출", "hanja": "引出·끌인·날출", "definition": "예금 따위를 찾음.", "example": "현금을 {}하려고 은행에 갔다."},
            {"term": "수수료", "hanja": "手數料·손수·셈수·료", "definition": "어떤 일을 처리해 준 대가로 받는 요금.", "example": "다른 은행에서 돈을 찾으면 {}가 붙는다."},
            {"term": "연체", "hanja": "延滯·늘일연·막힐체", "definition": "정해진 기한 안에 돈을 내지 못하고 지체함.", "example": "요금을 {}하면 이자가 붙는다."},
            {"term": "할부", "hanja": "割賦·나눌할·부세부", "definition": "돈을 여러 번으로 나누어 냄.", "example": "휴대폰 요금을 12개월 {}로 냈다."},
            {"term": "대출", "hanja": "貸出·빌릴대·날출", "definition": "돈이나 물건 따위를 빌려주거나 빌림.", "example": "은행에서 학자금 {}을 받았다."},
            {"term": "신용", "hanja": "信用·믿을신·쓸용", "definition": "채무를 갚을 능력이 있다고 믿을 만한 상태.", "example": "대금을 제때 갚아 {}을 지켰다."},
        ],
        "checkpoints": [
            {"options": ("인출", "수수료"), "template": "다른 은행 현금인출기에서 돈을 {0}했더니 {1}가 붙어 있었다."},
            {"options": ("연체", "할부"), "template": "휴대폰 요금을 3개월 {1}로 냈는데, 한 번 {0}되어 이자가 붙었다."},
            {"options": ("대출", "신용"), "template": "제때 갚지 않으면 {1}이 나빠져 나중에 {0}받기 어려워진다."},
        ],
        "confusable": {
            "left": {"term": "예금", "hanja": "預金", "definition": "돈을 금융 기관에 맡김.", "example": "목돈이 생겨 한번에 예금했다."},
            "right": {"term": "적금", "hanja": "積金", "definition": "일정 금액을 정기적으로 저축하여 목돈을 만드는 저축.", "example": "매달 5만 원씩 적금을 붓는다."},
            "tip": "한번에 맡기면 예금, 매달 꾸준히 부으면 적금!",
        },
        "quote": {"text": "외상이면 소도 잡아먹는다.", "author": "한국 속담"},
    },

    "d3": {
        "tag": "확인①", "intro": "이틀 동안 배운 12개 어휘를 세 가지 방식으로 확인해 봅시다.",
        "sections": [
            {"title": "① 초성 힌트 — 빈칸에 낱말을 쓰세요", "type": "fill", "items": [
                {"prompt": "[ㅇㅅ] 비용을 미리 헤아려 계산함", "hint": "이번 달 용돈 ___을 세웠다."},
                {"prompt": "[ㅈㅊ] 절약하여 모아 둠", "hint": "매달 용돈을 ___한다."},
                {"prompt": "[ㅈㅇ] 쓰고 남은 나머지 돈", "hint": "통장 ___을 확인했다."},
                {"prompt": "[ㅇㅊ] 예금 따위를 찾음", "hint": "현금을 ___했다."},
            ]},
            {"title": "② 글자 카드 조합 — 지·출·이·자·입·금·신·용", "type": "fill", "items": [
                {"prompt": "어떤 목적을 위하여 돈을 씀"},
                {"prompt": "돈을 빌려 쓴 대가로 치르는 돈"},
                {"prompt": "계좌에 돈을 넣음"},
                {"prompt": "채무를 갚을 능력이 있다고 믿을 만한 상태"},
            ]},
            {"title": "③ 문장에 어울리는 어휘 고르기", "type": "choice", "items": [
                {"prompt": "이번 달은 교통비 (　) 이 늘었다.", "options": ["지출", "저축"]},
                {"prompt": "예금을 하면 (　) 가 붙는다.", "options": ["이자", "수수료"]},
                {"prompt": "다른 은행 기기에서 돈을 찾으면 (　) 가 붙는다.", "options": ["수수료", "연체"]},
                {"prompt": "요금을 제때 못 내면 (　) 가 된다.", "options": ["연체", "할부"]},
            ]},
        ],
        "quote": {"text": "저축한 한 푼은 번 한 푼과 같다.", "author": "벤저민 프랭클린"},
    },

    "d4": {
        "tag": "확인②", "intro": "밑줄 친 어휘의 뜻을 확인하고, 실제 안내문 속에서 어떻게 쓰이는지 살펴봅시다.",
        "sections": [
            {"title": "④ 밑줄 뜻으로 알맞은 것 고르기", "type": "choice4", "items": [
                {"prompt": "휴대폰 요금을 12개월 할부로 냈다.", "options": ["나누어 냄", "한번에 냄", "미리 냄", "늦게 냄"]},
                {"prompt": "은행에서 학자금 대출을 받았다.", "options": ["빌림", "저축함", "기부함", "투자함"]},
                {"prompt": "대금을 제때 갚아 신용을 지켰다.", "options": ["믿을 만한 상태", "빚진 상태", "남은 금액", "이자율"]},
                {"prompt": "통장 잔액을 확인했다.", "options": ["남은 금액", "빌린 금액", "이자율", "수수료"]},
            ]},
        ],
        "passage": {
            "title": "⑤ 안내문을 읽고 답하세요",
            "body": [
                "[청소년 자유적금 안내]",
                "1. 대상: 만 14~19세 청소년   2. 금액: 매달 1만~20만 원 자유 입금",
                "3. 이자: 기본 금리 연 3% + 우대금리 최대 연 2%",
                "4. 유의: 정해진 날짜에 납입하지 않아도 되나, 중도 해지 시 이자가 줄어듦",
                "5. 인출: 만기 전 제한, 급할 때는 예금 잔액 일부만 대출 가능",
                "6. 문의: 가까운 은행 영업점 또는 상담 창구 방문",
            ],
            "items": [
                {"prompt": "이 적금에 가입할 수 있는 사람은?", "options": ["만 14~19세 청소년", "만 20세 이상 성인", "초등학생", "제한 없음"]},
                {"prompt": "이 적금의 이자 설명으로 알맞은 것은?", "options": ["기본 금리+우대금리", "이자 없음", "매달 이자율 변경", "가입 즉시 지급"]},
                {"prompt": "중도에 해지하면?", "options": ["이자가 줄어든다", "원금을 잃는다", "벌금을 낸다", "변화 없다"]},
                {"prompt": "만기 전 급히 필요할 때는?", "options": ["잔액 일부로 대출받기", "전액 인출하기", "양도하기", "이자만 받기"]},
                {"prompt": "안내문의 목적은?", "options": ["적금 상품 안내", "은행 위치 안내", "세금 안내", "취업 안내"]},
            ],
        },
        "quote": {"text": "가격은 당신이 지불하는 것이고, 가치는 당신이 얻는 것이다.", "author": "워런 버핏"},
    },

    "d5": {
        "tag": "종합복습", "intro": "이번 주에 배운 12개 어휘를 새로운 문장으로 마지막으로 점검해 봅시다.",
        "sections": [
            {"title": "뜻을 보고 어휘를 쓰세요", "type": "fill", "items": [
                {"prompt": "어떤 목적을 위하여 돈을 씀"},
                {"prompt": "절약하여 모아 둠"},
                {"prompt": "쓰고 남은 나머지 돈"},
                {"prompt": "예금 따위를 찾음"},
                {"prompt": "정해진 기한 안에 돈을 내지 못하고 지체함"},
                {"prompt": "채무를 갚을 능력이 있다고 믿을 만한 상태"},
            ]},
            {"title": "문장을 읽고 알맞은 어휘를 고르세요", "type": "choice", "items": [
                {"prompt": "이번 달 용돈 (　)을 세웠다.", "options": ["예산", "수수료"]},
                {"prompt": "예금을 하면 (　)가 붙는다.", "options": ["이자", "할부"]},
                {"prompt": "통장에 세뱃돈이 (　)되었다.", "options": ["입금", "연체"]},
                {"prompt": "다른 은행에서 돈을 찾으면 (　)가 붙는다.", "options": ["수수료", "신용"]},
                {"prompt": "휴대폰 요금을 3개월 (　)로 냈다.", "options": ["할부", "대출"]},
                {"prompt": "은행에서 학자금 (　)을 받았다.", "options": ["대출", "저축"]},
            ]},
        ],
        "quote": {"text": "돈을 쓰기 전에 먼저 저축하라. 저축하고 남은 돈을 써라.", "author": "워런 버핏"},
    },
}

WEEK3 = {
    "week": 3, "unit": "Ⅰ. 생존 문해력", "topic": "소비자·계약(약관) 읽기",

    "d1": {
        "tag": "학습①", "intro": "온라인 쇼핑, 휴대폰 개통, 아르바이트 계약서에는 낯선 소비자·계약 용어가 자주 나와요. 6개 어휘부터 익혀 봅시다.",
        "words": [
            {"term": "약관", "hanja": "約款·맺을약·항목관", "definition": "계약 당사자가 지켜야 할 사항을 미리 정해 놓은 조항.", "example": "가입 전에 이용 {}을 꼼꼼히 읽어야 한다."},
            {"term": "계약", "hanja": "契約·맺을계·맺을약", "definition": "서로 지켜야 할 의무를 글이나 말로 정하여 둠.", "example": "아르바이트를 시작하기 전에 근로 {}을 맺었다."},
            {"term": "배송", "hanja": "配送·나눌배·보낼송", "definition": "물건을 보내 줌.", "example": "주문한 물건의 {}이 하루 늦어졌다."},
            {"term": "하자", "hanja": "瑕疵·티하·흠자", "definition": "흠이나 잘못. 정상적인 상태를 갖추지 못한 결점.", "example": "새로 산 신발에서 {}를 발견했다."},
            {"term": "반품", "hanja": "返品·돌이킬반·물건품", "definition": "사들인 물건을 도로 돌려보냄.", "example": "사이즈가 맞지 않아 {}을 신청했다."},
            {"term": "교환", "hanja": "交換·사귈교·바꿀환", "definition": "서로 바꿈.", "example": "색상이 마음에 안 들어 다른 색으로 {}했다."},
        ],
        "checkpoints": [
            {"options": ("약관", "계약"), "template": "물건을 사기 전에 이용 {0}을 꼼꼼히 읽어야 하고, 아르바이트를 할 때는 근로 {1}서를 반드시 작성해야 한다."},
            {"options": ("배송", "하자"), "template": "주문한 물건의 {0}이 늦어지더니, 막상 받아 보니 제품에 {1}까지 있었다."},
            {"options": ("반품", "교환"), "template": "사이즈가 맞지 않아 {0}을 신청했는데, 매장에서는 환불 대신 다른 사이즈로 {1}해 주겠다고 했다."},
        ],
        "quote": {"text": "악마는 디테일에 있다.", "author": "서양 격언"},
    },

    "d2": {
        "tag": "학습②", "intro": "어제 배운 6개에 이어, 나머지 6개 어휘를 익혀 봅시다.",
        "words": [
            {"term": "청약", "hanja": "請約·청할청·맺을약", "definition": "계약을 맺자는 의사 표시.", "example": "인터넷으로 상품 구매 {}을 했다."},
            {"term": "철회", "hanja": "撤回·거둘철·돌아올회", "definition": "이미 낸 의사 표시를 거두어들이거나 취소함.", "example": "마음이 바뀌어 주문을 {}했다."},
            {"term": "위약금", "hanja": "違約金·어길위·맺을약·돈금", "definition": "계약을 지키지 않을 경우 벌로 물어야 하는 돈.", "example": "계약을 중도에 해지하면 {}을 내야 한다."},
            {"term": "면책", "hanja": "免責·면할면·꾸짖을책", "definition": "책임이나 책망을 면함.", "example": "천재지변으로 인한 피해는 {} 조항에 해당한다."},
            {"term": "손해배상", "hanja": "損害賠償·덜손·물어줄배", "definition": "남에게 끼친 손해를 물어 줌.", "example": "물건이 파손되어 {}을 청구했다."},
            {"term": "소비자", "hanja": "消費者·쓸비·사람자", "definition": "재화나 용역을 돈을 주고 사서 쓰는 사람.", "example": "{}의 권리를 보호하는 법이 있다."},
        ],
        "checkpoints": [
            {"options": ("청약", "철회"), "template": "인터넷으로 상품 구매 {0}을 했다가, 마음이 바뀌어 다음 날 바로 {1}했다."},
            {"options": ("위약금", "면책"), "template": "계약을 중도에 해지하면 {0}을 내야 하지만, 천재지변 등은 {1} 조항에 해당해 예외로 인정된다."},
            {"options": ("손해배상", "소비자"), "template": "택배로 온 물건이 파손되어 {0}을 청구했고, {1} 보호 센터에 상담도 신청했다."},
        ],
        "confusable": {
            "left": {"term": "환불", "hanja": "還拂", "definition": "이미 지불한 돈을 그대로 되돌려 줌.", "example": "불량품이라 환불받았다."},
            "right": {"term": "환급", "hanja": "還給", "definition": "이미 낸 세금이나 요금 중 일부를 돌려줌.", "example": "초과 납부한 세금을 환급받았다."},
            "tip": "낸 돈 전부를 돌려받으면 환불, 일부만 돌려받으면 환급!",
        },
        "quote": {"text": "돌다리도 두들겨 보고 건너라.", "author": "한국 속담"},
    },

    "d3": {
        "tag": "확인①", "intro": "이틀 동안 배운 12개 어휘를 세 가지 방식으로 확인해 봅시다.",
        "sections": [
            {"title": "① 초성 힌트 — 빈칸에 낱말을 쓰세요", "type": "fill", "items": [
                {"prompt": "[ㅇㄱ] 지켜야 할 사항을 미리 정한 조항", "hint": "이용 ___을 꼼꼼히 읽는다."},
                {"prompt": "[ㄱㅇ] 서로 지켜야 할 의무를 정하여 둠", "hint": "근로 ___서를 작성했다."},
                {"prompt": "[ㅂㅅ] 물건을 보내 줌", "hint": "물건의 ___이 늦어졌다."},
                {"prompt": "[ㅎㅈ] 흠이나 잘못", "hint": "제품에서 ___를 발견했다."},
            ]},
            {"title": "② 글자 카드 조합 — 반·품·교·환·청·약·철·회", "type": "fill", "items": [
                {"prompt": "사들인 물건을 도로 돌려보냄"},
                {"prompt": "서로 바꿈"},
                {"prompt": "계약을 맺자는 의사 표시"},
                {"prompt": "이미 낸 의사 표시를 거두어들이거나 취소함"},
            ]},
            {"title": "③ 문장에 어울리는 어휘 고르기", "type": "choice", "items": [
                {"prompt": "사이즈가 안 맞아 (　) 을 신청했다.", "options": ["반품", "교환"]},
                {"prompt": "색상이 안 들어 다른 색으로 (　) 했다.", "options": ["교환", "철회"]},
                {"prompt": "인터넷으로 상품 구매 (　) 을 했다.", "options": ["청약", "위약금"]},
                {"prompt": "마음이 바뀌어 주문을 (　) 했다.", "options": ["철회", "면책"]},
            ]},
        ],
        "quote": {"text": "아는 것이 힘이다.", "author": "프랜시스 베이컨"},
    },

    "d4": {
        "tag": "확인②", "intro": "밑줄 친 어휘의 뜻을 확인하고, 실제 쇼핑몰 약관 속에서 어떻게 쓰이는지 살펴봅시다.",
        "sections": [
            {"title": "④ 밑줄 뜻으로 알맞은 것 고르기", "type": "choice4", "items": [
                {"prompt": "계약을 중도에 해지하면 위약금을 내야 한다.", "options": ["벌로 무는 돈", "빌려주는 돈", "돌려받는 돈", "이자로 붙는 돈"]},
                {"prompt": "천재지변으로 인한 피해는 면책 조항에 해당한다.", "options": ["책임을 면함", "책임을 더함", "손해를 물어줌", "계약을 취소함"]},
                {"prompt": "물건이 파손되어 손해배상을 청구했다.", "options": ["끼친 손해를 물어 줌", "물건을 교환함", "세금을 돌려받음", "계약을 다시 맺음"]},
                {"prompt": "소비자의 권리를 보호하는 법이 있다.", "options": ["재화·용역을 사서 쓰는 사람", "물건을 만드는 사람", "물건을 파는 사람", "광고를 하는 사람"]},
            ]},
        ],
        "passage": {
            "title": "⑤ 약관을 읽고 답하세요",
            "body": [
                "[온라인 쇼핑몰 이용약관 (발췌)]",
                "1. 청약 철회: 상품을 받은 날부터 7일 이내 철회 가능",
                "2. 반품·교환: 하자 있는 경우, 배송받은 날로부터 30일 이내 신청 가능",
                "3. 환불: 단순 변심 반품은 왕복 배송비를 소비자가 부담",
                "4. 위약금: 예약 상품을 정당한 사유 없이 취소 시 부과될 수 있음",
                "5. 면책: 천재지변 등 불가항력으로 인한 배송 지연은 회사 책임 없음",
                "6. 문의: 고객센터를 통해 상담 가능",
            ],
            "items": [
                {"prompt": "청약을 철회할 수 있는 기간은?", "options": ["받은 날부터 7일 이내", "주문 당일만", "30일 이내 언제든", "기간 제한 없음"]},
                {"prompt": "하자 제품의 반품·교환 신청 기간은?", "options": ["배송받은 날로부터 30일", "7일 이내", "1년 이내", "신청 불가"]},
                {"prompt": "단순 변심 반품 시 배송비 부담은?", "options": ["소비자", "회사", "반반 부담", "배송업체"]},
                {"prompt": "위약금이 부과될 수 있는 경우는?", "options": ["정당한 사유 없이 예약 취소", "제품 하자", "배송 지연", "청약 철회"]},
                {"prompt": "약관 내용과 일치하지 않는 것은?", "options": ["천재지변 배송 지연도 회사 책임", "하자 제품 30일 이내 반품", "단순 변심 반품비 소비자 부담", "청약 7일 이내 철회 가능"]},
            ],
        },
        "quote": {"text": "이행할 수 있는 것보다 더 많이 약속하지 마라.", "author": "벤저민 프랭클린"},
    },

    "d5": {
        "tag": "종합복습", "intro": "이번 주에 배운 12개 어휘를 새로운 문장으로 마지막으로 점검해 봅시다.",
        "sections": [
            {"title": "뜻을 보고 어휘를 쓰세요", "type": "fill", "items": [
                {"prompt": "지켜야 할 사항을 미리 정해 놓은 조항"},
                {"prompt": "물건을 보내 줌"},
                {"prompt": "흠이나 잘못"},
                {"prompt": "이미 낸 의사 표시를 거두어들이거나 취소함"},
                {"prompt": "계약을 지키지 않을 경우 벌로 물어야 하는 돈"},
                {"prompt": "재화나 용역을 사서 쓰는 사람"},
            ]},
            {"title": "문장을 읽고 알맞은 어휘를 고르세요", "type": "choice", "items": [
                {"prompt": "이용 (　)을 꼼꼼히 읽어야 한다.", "options": ["약관", "배송"]},
                {"prompt": "주문한 물건의 (　)이 늦어졌다.", "options": ["배송", "하자"]},
                {"prompt": "제품에서 (　)를 발견했다.", "options": ["하자", "반품"]},
                {"prompt": "마음이 바뀌어 주문을 (　)했다.", "options": ["철회", "교환"]},
                {"prompt": "계약을 중도 해지하면 (　)을 내야 한다.", "options": ["위약금", "면책"]},
                {"prompt": "(　)의 권리를 보호하는 법이 있다.", "options": ["소비자", "청약"]},
            ]},
        ],
        "quote": {"text": "정직이 최선의 방책이다.", "author": "서양 속담"},
    },
}

CIRCLE_NUM = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩", "⑪", "⑫"]
CHOICE_MARK = ["㉮", "㉯", "㉰", "㉱"]


# =========================================================================
# 저수준 헬퍼
# =========================================================================

def keep_row_together(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def keep_table_together(table):
    for row in table.rows:
        keep_row_together(row)


def merge_last_row_if_odd(table, count):
    """2단 표에서 항목 수가 홀수면 마지막 행의 빈 오른쪽 칸을 왼쪽 칸과 합쳐 없앤다."""
    if count % 2 == 1:
        last_row = table.rows[-1]
        last_row.cells[0].merge(last_row.cells[1])


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def set_run_font(run, size=9.5, bold=False, color=None, name=FONT_NAME):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def tight(paragraph, before=0, after=2, line=1.0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    return paragraph


def p(doc_or_cell, text="", size=9.5, bold=False, color=None, align=None, before=0, after=2):
    para = doc_or_cell.add_paragraph()
    tight(para, before, after)
    if align is not None:
        para.alignment = align
    if text:
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return para


def mixed_run(paragraph, pieces):
    """pieces: [(text, size, bold, color), ...]"""
    for text, size, bold, color in pieces:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)


def section_bar(doc, text, subtitle=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "3D4BC7")
    set_cell_margins(cell, top=45, bottom=45, left=120, right=120)
    para = cell.paragraphs[0]
    tight(para, 0, 0)
    run = para.add_run(text)
    set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    if subtitle:
        p2 = p(cell, subtitle, size=7.8, color=RGBColor(0xDD, 0xE0, 0xFF), after=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def info_header(doc, week, topic, day_label, intro):
    p(doc, "명품 경북여상 문해력 15분  ·  손으로 쓰는 학습지", size=14, bold=True, color=ACCENT, after=1)

    sub = doc.add_paragraph()
    tight(sub, 0, 3)
    sub.paragraph_format.tab_stops.add_tab_stop(Cm(17.8), WD_TAB_ALIGNMENT.RIGHT)
    r1 = sub.add_run(f"{week}주차 · {topic}  —  {day_label}")
    set_run_font(r1, size=10.5, bold=True)
    sub.add_run("\t")
    r2 = sub.add_run("학습일   ")
    set_run_font(r2, size=9.5, bold=True, color=GRAY)
    r3 = sub.add_run(f"{BLANK[:3]} 월  {BLANK[:3]} 일")
    set_run_font(r3, size=11.5)

    p(doc, intro, size=8.5, color=GRAY, before=2, after=4)


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)


# ---- Day1: 단어 카드 (2단 배치) ----

def word_cell_content(cell, idx, w):
    set_cell_margins(cell, top=35, bottom=45, left=100, right=100)
    head = tight(cell.paragraphs[0], 1, 0)
    mixed_run(head, [
        (f"{CIRCLE_NUM[idx]} {w['term']}  ", 12, True, None),
        (f"[{w['hanja']}]", 7.5, False, GRAY),
    ])
    d = p(cell, after=0)
    mixed_run(d, [("뜻  ", 8, True, ACCENT), (w["definition"], 8.5, False, None)])
    ex = p(cell, after=1)
    mixed_run(ex, [("예문  ", 8, True, ACCENT), (w["example"].format(BLANK), 8.5, False, None)])

    inner = cell.add_table(rows=2, cols=1)
    inner.rows[0].cells[0].width = Cm(8.2)
    inner.rows[1].cells[0].width = Cm(8.2)
    label_cell = inner.rows[0].cells[0]
    set_cell_shading(label_cell, LIGHTGRAY)
    set_cell_margins(label_cell, 15, 15, 60, 60)
    lp = tight(label_cell.paragraphs[0], 0, 0)
    lr = lp.add_run("✍ 나의 글씨체로 예문 완성하기")
    set_run_font(lr, size=7.3, bold=True, color=GRAY)

    write_cell = inner.rows[1].cells[0]
    set_cell_margins(write_cell, 15, 15, 60, 60)
    tr = inner.rows[1]._tr
    trPr = tr.get_or_add_trPr()
    trh = OxmlElement("w:trHeight")
    trh.set(qn("w:val"), "170")
    trh.set(qn("w:hRule"), "atLeast")
    trPr.append(trh)


def word_grid(doc, words):
    rows = (len(words) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    keep_table_together(table)
    for i, w in enumerate(words):
        r, c = divmod(i, 2)
        cell = table.rows[r].cells[c]
        cell.width = Cm(8.6)
        word_cell_content(cell, i, w)
    merge_last_row_if_odd(table, len(words))


def checkpoint_grid(doc, checkpoints):
    rows = (len(checkpoints) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    keep_table_together(table)
    for i, cpb in enumerate(checkpoints):
        r, c = divmod(i, 2)
        cell = table.rows[r].cells[c]
        cell.width = Cm(8.6)
        set_cell_margins(cell, 40, 45, 90, 90)
        sent = cpb["template"].format(f"{BLANK}", f"{BLANK}")
        head = tight(cell.paragraphs[0], 1, 1)
        mixed_run(head, [(f"Q{i+1}. ", 9, True, ACCENT), (sent, 9, False, None)])
        hint = p(cell, f"보기: {cpb['options'][0]} · {cpb['options'][1]}", size=7.5, color=GRAY, after=1)
    merge_last_row_if_odd(table, len(checkpoints))


def confusable_box(doc, c):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    keep_table_together(table)
    for cell, side in zip(table.rows[0].cells, (c["left"], c["right"])):
        cell.width = Cm(8.6)
        set_cell_margins(cell, 40, 40, 90, 90)
        set_cell_shading(cell, "F7F8FF")
        head = tight(cell.paragraphs[0], 0, 1)
        mixed_run(head, [(f"{side['term']} ", 10.5, True, ACCENT), (f"[{side['hanja']}]  ", 7.5, False, GRAY),
                          (side["definition"], 8.5, False, None)])
        ex = p(cell, "예) " + side["example"], size=8, color=GRAY, after=0)
    tip = p(doc, "🧠 " + c["tip"], size=8.5, color=ACCENT, before=2, after=1)


# ---- Day2: 문제 ----

def fill_grid(doc, items):
    rows = (len(items) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    keep_table_together(table)
    for i, item in enumerate(items):
        r, c = divmod(i, 2)
        cell = table.rows[r].cells[c]
        cell.width = Cm(8.6)
        set_cell_margins(cell, 35, 40, 90, 90)
        head = tight(cell.paragraphs[0], 1, 1)
        mixed_run(head, [(f"{i+1}. ", 9, True, None), (item["prompt"], 9, False, None)])
        if item.get("hint"):
            p(cell, item["hint"].replace("___", BLANK[:4]), size=8, color=GRAY, after=1)
        ans = p(cell, after=1)
        mixed_run(ans, [("답  ", 8, True, ACCENT), (BLANK, 9, False, None)])
    merge_last_row_if_odd(table, len(items))


def choice_grid(doc, items, four=False):
    rows = (len(items) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    keep_table_together(table)
    marks = CHOICE_MARK if four else ["㉮", "㉯"]
    for i, item in enumerate(items):
        r, c = divmod(i, 2)
        cell = table.rows[r].cells[c]
        cell.width = Cm(8.6)
        set_cell_margins(cell, 35, 40, 90, 90)
        head = tight(cell.paragraphs[0], 1, 1)
        mixed_run(head, [(f"{i+1}. ", 9, True, None), (item["prompt"], 9, False, None)])
        opts = " ".join(f"{marks[j]}{o}" for j, o in enumerate(item["options"]))
        p(cell, opts, size=8, color=GRAY, after=1)
        ans = p(cell, after=1)
        mixed_run(ans, [("답  ", 8, True, ACCENT), (BLANK[:5], 9, False, None)])
    merge_last_row_if_odd(table, len(items))


def passage_box(doc, passage):
    p(doc, passage["title"], size=10.5, bold=True, color=ACCENT, before=4, after=2)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    keep_table_together(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHTGRAY)
    set_cell_margins(cell, 50, 50, 100, 100)
    first = True
    for line in passage["body"]:
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        tight(para, 0, 1)
        run = para.add_run(line)
        set_run_font(run, size=8.3, bold=line.startswith("["))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    for i, item in enumerate(passage["items"]):
        para = p(doc, after=0)
        mixed_run(para, [(f"{i+1}. ", 9, True, None), (item["prompt"] + "  ", 9, False, None)])
        opts_run = para.add_run(" ".join(f"{CHOICE_MARK[j]}{o}" for j, o in enumerate(item["options"])))
        set_run_font(opts_run, size=7.8, color=GRAY)
        ans = p(doc, after=2)
        mixed_run(ans, [("     답  ", 8, True, ACCENT), (BLANK[:5], 9, False, None)])


def reflection_block(doc):
    section_bar(doc, "오늘의 한 줄", "가장 기억에 남는 낱말과 이유를 써 보세요.")
    ruled_line(doc, before=2, after=4)


def set_cell_border_bottom_only(cell, size=6, color="9AA0C3"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    tcPr.append(borders)


def writing_lines(doc, count=2, row_height=300):
    """서로 분리된 줄로 확실히 보이는 밑줄(표 기반). 문단 테두리는 인접하면 하나로 합쳐 보이는 문제가 있어 표로 대체."""
    table = doc.add_table(rows=count, cols=1)
    for row in table.rows:
        cell = row.cells[0]
        cell.width = Cm(17.8)
        set_cell_margins(cell, 40, 40, 20, 20)
        set_cell_border_bottom_only(cell)
        trPr = row._tr.get_or_add_trPr()
        trh = OxmlElement("w:trHeight")
        trh.set(qn("w:val"), str(row_height))
        trh.set(qn("w:hRule"), "atLeast")
        trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def ruled_line(doc, before=4, after=10):
    para = doc.add_paragraph()
    tight(para, before, after)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "9AA0C3")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def quote_box(doc, quote):
    p(doc, "💬 오늘의 명언", size=9.5, bold=True, color=ACCENT, before=4, after=2)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    keep_table_together(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F7F8FF")
    set_cell_margins(cell, 35, 35, 120, 120)
    qp = cell.paragraphs[0]
    tight(qp, 0, 1)
    qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = qp.add_run(f"“{quote['text']}”")
    run.font.italic = True
    set_run_font(run, size=10, bold=True, color=RGBColor(0x33, 0x33, 0x66))
    run.font.italic = True
    ap = cell.add_paragraph()
    tight(ap, 0, 0)
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    arun = ap.add_run(f"— {quote['author']}")
    set_run_font(arun, size=8, color=GRAY)

    guide = p(doc, "✍ 위 명언을 소리 내어 한 번 읽고, 나만의 글씨체로 또박또박 따라 써 보세요.",
              size=8, color=GRAY, before=2, after=1)
    writing_lines(doc, count=2, row_height=160)


def todo_box(doc, count=3):
    p(doc, "✅ 오늘 해야 할 일", size=9.5, bold=True, color=ACCENT, before=0, after=2)
    table = doc.add_table(rows=1, cols=count)
    table.style = "Table Grid"
    keep_table_together(table)
    col_w = 17.8 / count
    for i in range(count):
        cell = table.rows[0].cells[i]
        cell.width = Cm(col_w)
        set_cell_margins(cell, 35, 35, 80, 80)
        para = cell.paragraphs[0]
        tight(para, 0, 0)
        run = para.add_run("☐ ")
        set_run_font(run, size=10, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def base_document():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(9.5)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    return doc


def render_sections(doc, sections):
    for sec in sections:
        p(doc, sec["title"], size=9.5, bold=True, color=ACCENT, before=3, after=1)
        if sec["type"] == "fill":
            fill_grid(doc, sec["items"])
        elif sec["type"] == "choice":
            choice_grid(doc, sec["items"], four=False)
        elif sec["type"] == "choice4":
            choice_grid(doc, sec["items"], four=True)


def build_day(topic, day_key, day, out_path):
    doc = base_document()
    info_header(doc, topic["week"], topic["topic"], f"Day {day_key[1]} · {day['tag']}", day["intro"])
    todo_box(doc, count=3)

    if "words" in day:  # 학습 day
        word_grid(doc, day["words"])
        if day.get("checkpoints"):
            p(doc, "확인 문제 — 알맞은 말을 빈칸에 쓰세요", size=9, bold=True, color=ACCENT, before=1, after=1)
            checkpoint_grid(doc, day["checkpoints"])
        if day.get("confusable"):
            p(doc, "⭐ 헷갈리기 쉬운 어휘", size=9, bold=True, color=ACCENT, before=1, after=1)
            confusable_box(doc, day["confusable"])
    else:  # 확인/복습 day
        render_sections(doc, day.get("sections", []))
        if day.get("passage"):
            passage_box(doc, day["passage"])
        if day_key == "d5":
            reflection_block(doc)

    if day.get("quote"):
        quote_box(doc, day["quote"])

    save_with_retry(doc, out_path)
    print("saved", out_path)


def save_with_retry(doc, out_path, attempts=6, delay=1.0):
    """OneDrive가 방금 만든 파일을 잠깐 동기화하며 잠그는 경우가 있어 재시도한다."""
    import time
    for i in range(attempts):
        try:
            doc.save(out_path)
            return
        except PermissionError:
            if i == attempts - 1:
                raise
            time.sleep(delay)


def build_week(topic, out_prefix):
    for day_key in ["d1", "d2", "d3", "d4", "d5"]:
        day = topic[day_key]
        build_day(topic, day_key, day, f"{out_prefix}_Day{day_key[1]}_학습지.docx")


if __name__ == "__main__":
    build_week(WEEK1, "1주차_1_공문서행정문서읽기")
    build_week(WEEK2, "2주차_2_금융경제생활")
    build_week(WEEK3, "3주차_3_소비자계약약관읽기")
