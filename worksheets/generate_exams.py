# -*- coding: utf-8 -*-
"""
문해력 15분 — 정기고사(중간·기말) 문제지 생성기

중간고사·기말고사 각 2회(1학기/2학기), 회차별 20문항 4지선다형, 문항당 1점(총 20점).
문서 하단에 정답표와 문항별 간단 해설을 포함한다.
출력 위치: worksheets/정기고사/

사용법:
    python generate_exams.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FONT_NAME = "맑은 고딕"
ACCENT = RGBColor(0x3D, 0x4B, 0xC7)
GRAY = RGBColor(0x6B, 0x6F, 0x85)
LIGHTGRAY = "F2F3F8"
CIRCLE = ["①", "②", "③", "④"]

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "정기고사")

# =========================================================================
# 저수준 헬퍼 (worksheets/generate.py 와 동일한 스타일)
# =========================================================================

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_run_font(run, size=10.5, bold=False, color=None, name=FONT_NAME):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def tight(paragraph, before=0, after=4, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def p(doc_or_cell, text="", size=10.5, bold=False, color=None, align=None, before=0, after=4):
    para = doc_or_cell.add_paragraph()
    if align:
        para.alignment = align
    tight(para, before, after)
    if text:
        run = para.add_run(text)
        set_run_font(run, size, bold, color)
    return para


def mixed_run(paragraph, pieces):
    for text, size, bold, color in pieces:
        run = paragraph.add_run(text)
        set_run_font(run, size, bold, color)


def base_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    return doc


def add_page_break(doc):
    doc.add_page_break()


def exam_header(doc, exam_title, exam_range):
    title_p = doc.add_paragraph()
    tight(title_p, 0, 2)
    mixed_run(title_p, [("명품 경북여상 문해력 15분", 15, True, ACCENT), ("  ·  정기고사", 13, True, GRAY)])

    sub_p = doc.add_paragraph()
    tight(sub_p, 0, 6)
    mixed_run(sub_p, [(exam_title, 13.5, True, None)])

    # 응시자 정보 입력란
    info = doc.add_table(rows=1, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Cm(3.0), Cm(3.0), Cm(3.5), Cm(5.5)]
    labels = ["학년 ___ 반", "번호 ___ 번", "이름", "________________"]
    for i, cell in enumerate(info.rows[0].cells):
        cell.width = widths[i]
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        cp = cell.paragraphs[0]
        tight(cp, 0, 0)
        run = cp.add_run(labels[i])
        set_run_font(run, 10, False, GRAY)

    note = doc.add_paragraph()
    tight(note, 6, 2)
    mixed_run(note, [
        (f"시험 범위: {exam_range}", 9.5, False, GRAY),
    ])
    note2 = doc.add_paragraph()
    tight(note2, 0, 10)
    mixed_run(note2, [
        ("배점 안내: 전 문항 각 1점, 총 20점 만점 · 4지선다형 · 가장 알맞은 답을 하나만 고르세요.", 9.5, True, ACCENT),
    ])
    ruled = doc.add_paragraph()
    tight(ruled, 0, 8)
    pPr = ruled._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "3D4BC7")
    pbdr.append(bottom)
    pPr.append(pbdr)


def render_question(doc, qnum, item):
    if item.get("passage_intro"):
        box = doc.add_table(rows=1, cols=1)
        box.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = box.rows[0].cells[0]
        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        set_cell_shading(cell, LIGHTGRAY)
        cp = cell.paragraphs[0]
        tight(cp, 0, 0)
        run = cp.add_run(item["passage_intro"])
        set_run_font(run, 9.7, False, None)
        run.italic = False
        after_gap = doc.add_paragraph()
        tight(after_gap, 0, 4)

    qp = doc.add_paragraph()
    tight(qp, 4, 3)
    mixed_run(qp, [(f"{qnum}. ", 10.5, True, ACCENT), (item["prompt"], 10.5, False, None)])

    if item.get("underline"):
        # underline handled inline within prompt text already using <b> markup replaced upstream
        pass

    opt_table = doc.add_table(rows=2, cols=2)
    opt_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for c in opt_table.columns:
        c.width = Cm(8.9)
    idx = 0
    for r in range(2):
        for cidx in range(2):
            cell = opt_table.rows[r].cells[cidx]
            set_cell_margins(cell, top=30, bottom=30, left=60, right=60)
            cp = cell.paragraphs[0]
            tight(cp, 0, 0)
            run = cp.add_run(f"{CIRCLE[idx]} {item['options'][idx]}")
            set_run_font(run, 10, False, None)
            idx += 1
    gap = doc.add_paragraph()
    tight(gap, 0, 6)


def build_exam(exam_key, exam_title, exam_range, questions, out_prefix):
    doc = base_document()
    exam_header(doc, exam_title, exam_range)

    for i, item in enumerate(questions, start=1):
        render_question(doc, i, item)

    # 정답표
    add_page_break(doc)
    ans_title = doc.add_paragraph()
    tight(ans_title, 4, 6)
    mixed_run(ans_title, [("정답표", 13, True, ACCENT)])

    ans_table = doc.add_table(rows=4, cols=10)
    ans_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in ans_table.columns:
        c.width = Cm(1.65)
    for i, item in enumerate(questions):
        qn_ = i + 1
        col = (qn_ - 1) % 10
        row_no = ((qn_ - 1) // 10) * 2
        num_cell = ans_table.rows[row_no].cells[col]
        set_cell_margins(num_cell, top=30, bottom=20, left=20, right=20)
        set_cell_shading(num_cell, LIGHTGRAY)
        ncp = num_cell.paragraphs[0]
        ncp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tight(ncp, 0, 0)
        nrun = ncp.add_run(str(qn_))
        set_run_font(nrun, 9.5, True, GRAY)

        ans_cell = ans_table.rows[row_no + 1].cells[col]
        set_cell_margins(ans_cell, top=20, bottom=30, left=20, right=20)
        acp = ans_cell.paragraphs[0]
        acp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tight(acp, 0, 0)
        arun = acp.add_run(CIRCLE[item["answer"]])
        set_run_font(arun, 11, True, ACCENT)

    # 해설
    exp_title = doc.add_paragraph()
    tight(exp_title, 14, 6)
    mixed_run(exp_title, [("해설", 13, True, ACCENT)])

    for i, item in enumerate(questions, start=1):
        ep = doc.add_paragraph()
        tight(ep, 0, 3)
        mixed_run(ep, [
            (f"{i}. ", 9.3, True, ACCENT),
            (f"정답 {CIRCLE[item['answer']]}. ", 9.3, True, None),
            (item["explanation"], 9.3, False, GRAY),
        ])

    os.makedirs(OUT_DIR, exist_ok=True)
    out_path = os.path.join(OUT_DIR, f"{out_prefix}.docx")
    save_with_retry(doc, out_path)
    print("saved", out_path)


def save_with_retry(doc, out_path, attempts=6, delay=1.0):
    import time
    for i in range(attempts):
        try:
            doc.save(out_path)
            return
        except PermissionError:
            if i == attempts - 1:
                raise
            time.sleep(delay)


# =========================================================================
# 문항 데이터
# =========================================================================

EXAM1 = [
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 증명서 따위를 발행하여 내줌.",
     "options": ["발급", "신청", "열람", "통보"], "answer": 0,
     "explanation": "'발급'은 증명서 따위를 발행하여 내주는 것을 뜻한다. '신청'은 청구함, '열람'은 훑어봄, '통보'는 알림을 뜻한다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 제출한 문서를 처리하지 않고 되돌려 보냄.",
     "options": ["갱신", "사본", "반려", "공고"], "answer": 2,
     "explanation": "'반려'는 제출한 문서를 처리하지 않고 되돌려 보내는 것이다. '갱신'은 고쳐 새롭게 함, '사본'은 옮긴 문서, '공고'는 널리 알림을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 이번 달 용돈 ( )을 세워 보았다.",
     "options": ["지출", "예산", "잔액", "이자"], "answer": 1,
     "explanation": "필요한 비용을 미리 헤아려 계산하는 것은 '예산'이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 다른 은행 기기에서 돈을 찾으면 ( )가 붙는다.",
     "options": ["이자", "연체", "할부", "수수료"], "answer": 3,
     "explanation": "일을 처리해 준 대가로 받는 요금은 '수수료'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 계약을 맺자는 의사 표시.",
     "options": ["청약", "철회", "약관", "위약금"], "answer": 0,
     "explanation": "'청약'은 계약을 맺자는 의사 표시이다. '철회'는 의사 표시를 거두어들임을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 사이즈가 맞지 않아 ( )을 신청했다.",
     "options": ["교환", "하자", "반품", "배송"], "answer": 2,
     "explanation": "사들인 물건을 도로 돌려보내는 것은 '반품'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 사실과 다르게 해석하거나 그릇되게 함.",
     "options": ["편향", "왜곡", "조작", "검증"], "answer": 1,
     "explanation": "'왜곡'은 사실과 다르게 해석하거나 그릇되게 하는 것이다. '편향'은 한쪽으로 치우침을 뜻한다."},
    {"prompt": "밑줄 친 어휘의 뜻으로 알맞은 것은? — 이 자료는 <신뢰도>가 낮아 믿기 어렵다.",
     "options": ["사실과 다른 정도", "널리 퍼진 정도", "치우친 정도", "믿고 의지할 수 있는 정도"], "answer": 3,
     "explanation": "'신뢰도'는 믿고 의지할 수 있는 정도를 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 내 ( )에 맞는 진로를 찾고 있다.",
     "options": ["적성", "흥미", "경력", "자격증"], "answer": 0,
     "explanation": "어떤 일에 알맞은 성질이나 소질은 '적성'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 채용 시험 등에서 직접 만나 인품 등을 평가하는 시험.",
     "options": ["지원", "수습", "면접", "적응"], "answer": 2,
     "explanation": "'면접'은 직접 만나 평가하는 시험을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 매운맛 떡볶이의 ( )가 꾸준히 늘고 있다.",
     "options": ["공급", "수요", "매출", "손익"], "answer": 1,
     "explanation": "어떤 물건을 사려는 욕구는 '수요'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 상품을 다른 것과 구별하기 위한 고유의 이름이나 표시.",
     "options": ["마케팅", "광고", "홍보", "브랜드"], "answer": 3,
     "explanation": "'브랜드'는 상품을 구별하기 위한 고유의 이름이나 표시를 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 감기 때문에 병원에서 ( )를 받았다.",
     "options": ["진료", "처방", "진단", "접수"], "answer": 0,
     "explanation": "의사가 병을 진찰하고 치료하는 것은 '진료'이다."},
    {"prompt": "밑줄 친 어휘의 뜻으로 알맞은 것은? — 건강<보험> 덕분에 진료비 부담이 줄었다.",
     "options": ["진료 접수 순서", "약을 짓는 방법", "사고·질병에 대비해 보상받는 제도", "병의 진단 결과"], "answer": 2,
     "explanation": "'보험'은 사고·질병 등에 대비해 미리 돈을 내고 사고 시 보상받는 제도이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 병원체가 들어와 증식하는 일.",
     "options": ["격리", "감염", "완치", "병력"], "answer": 1,
     "explanation": "'감염'은 병원체가 들어와 증식하는 일이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 인물을 가운데 두는 ( )로 사진을 찍었다.",
     "options": ["여백", "채도", "명도", "구도"], "answer": 3,
     "explanation": "화면에 배치하는 짜임새는 '구도'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 저작물 등을 사용할 수 있도록 허락하는 것.",
     "options": ["라이선스", "초상권", "저작물", "편집"], "answer": 0,
     "explanation": "'라이선스'는 저작물 등의 사용을 허락하는 것을 뜻한다."},
    {"prompt": "밑줄 친 어휘의 뜻으로 알맞은 것은? — 친구가 영상에 나올 때는 <초상권> 동의를 받아야 한다.",
     "options": ["저작물을 사용할 권리", "영상을 편집할 권리", "얼굴 등이 함부로 촬영·공표되지 않을 권리", "음악을 만들 권리"], "answer": 2,
     "explanation": "'초상권'은 자신의 얼굴 등이 함부로 촬영·공표되지 않을 권리이다."},
    {"passage_intro": "[교내 진로 특강 신청 안내]\n1. 대상: 전교생 희망자\n2. 신청 기한: 이번 주 금요일까지 — 기한을 넘기면 접수되지 않습니다.\n3. 신청서 제출 후 담당 선생님이 접수 여부를 통보합니다.\n4. 서명이 빠진 신청서는 반려되니 유의하세요.\n5. 특강 자료는 특강 당일 배부되며, 이후 열람은 어렵습니다.",
     "prompt": "위 안내문에서, 신청 기한을 넘기면 어떻게 되는가?",
     "options": ["자동 연장된다", "접수되지 않는다", "벌점을 받는다", "다음 학기로 이월된다"], "answer": 1,
     "explanation": "안내문 2번에 기한을 넘기면 접수되지 않는다고 명시되어 있다."},
    {"passage_intro": None,
     "prompt": "위 안내문에서, 서명이 빠진 신청서는 어떻게 되는가?",
     "options": ["그대로 접수된다", "담당자가 대신 서명한다", "다음 특강에 자동 등록된다", "반려된다"], "answer": 3,
     "explanation": "안내문 4번에 서명이 빠진 신청서는 반려된다고 안내되어 있다."},
]

EXAM2 = [
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 생활이나 목적을 같이하는 집단.",
     "options": ["계층", "공동체", "소외", "인권"], "answer": 1,
     "explanation": "'공동체'는 생활이나 목적을 같이하는 집단을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 두 집단 사이에 ( )이 생겼다.",
     "options": ["연대", "배려", "봉사", "갈등"], "answer": 3,
     "explanation": "목표나 이해관계가 달라 서로 충돌하는 것은 '갈등'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 국민이 나라의 주인이 되어 권력을 행사하는 정치 제도.",
     "options": ["민주주의", "선거", "공약", "정책"], "answer": 0,
     "explanation": "'민주주의'는 국민이 나라의 주인이 되어 권력을 행사하는 정치 제도이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 국회는 후보자를 대상으로 ( )를 열었다.",
     "options": ["탄핵", "시행령", "청문회", "국회"], "answer": 2,
     "explanation": "이해관계인이나 전문가의 의견을 듣기 위해 여는 모임은 '청문회'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 법원에 재판을 청구함.",
     "options": ["판결", "소송", "항소", "합의"], "answer": 1,
     "explanation": "'소송'은 법원에 재판을 청구하는 것을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 검찰은 피의자를 ( )했다.",
     "options": ["항소", "합의", "배상", "기소"], "answer": 3,
     "explanation": "검사가 법원에 재판을 청구하는 것은 '기소'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 옛사람들이 남긴 흔적이 있는 곳.",
     "options": ["유적", "유물", "왕조", "전란"], "answer": 0,
     "explanation": "'유적'은 옛사람들이 남긴 흔적이 있는 곳이다."},
    {"prompt": "밑줄 친 어휘의 뜻으로 알맞은 것은? — 1945년 <광복>을 맞았다.",
     "options": ["하나였던 것이 둘로 나뉨", "전쟁으로 인한 어지러운 사태", "빼앗긴 나라의 주권을 도로 찾음", "제도를 새롭게 뜯어고침"], "answer": 2,
     "explanation": "'광복'은 빼앗긴 나라의 주권을 도로 찾는 것을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 과학자는 ( )을 세우고 실험으로 검증했다.",
     "options": ["관찰", "가설", "분석", "데이터"], "answer": 1,
     "explanation": "어떤 사실을 설명하기 위해 임시로 세운 이론은 '가설'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 다른 종류의 것이 녹아서 서로 구별이 없게 하나로 합침.",
     "options": ["혁신", "특허", "자동화", "융합"], "answer": 3,
     "explanation": "'융합'은 다른 종류의 것이 하나로 합쳐지는 것이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 지구 ( )로 빙하가 녹고 있다.",
     "options": ["온난화", "오염", "멸종", "탄소"], "answer": 0,
     "explanation": "지구의 평균 기온이 점점 올라가는 현상은 '온난화'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 계속 다시 만들어 쓸 수 있는 에너지.",
     "options": ["미세먼지", "대기", "재생에너지", "폐기물"], "answer": 2,
     "explanation": "'재생에너지'는 계속 다시 만들어 쓸 수 있는 에너지이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 미술관에서 작품을 ( )했다.",
     "options": ["전시", "감상", "공연", "창작"], "answer": 1,
     "explanation": "예술 작품을 음미하고 이해하는 것은 '감상'이다."},
    {"prompt": "밑줄 친 어휘의 뜻으로 알맞은 것은? — 자연을 <모티프>로 삼은 작품이다.",
     "options": ["사물의 좋고 나쁨을 논함", "창작의 계기가 되는 자극", "작품의 뜻을 풀이함", "작품에서 반복되는 중심 생각이나 이미지"], "answer": 3,
     "explanation": "'모티프'는 작품에서 반복되는 중심 생각이나 이미지를 뜻한다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 예로부터 전해 내려오는 사상·관습·행동 등의 양식.",
     "options": ["전통", "풍습", "유산", "세대"], "answer": 0,
     "explanation": "'전통'은 예로부터 전해 내려오는 사상·관습·행동 등의 양식이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 문화 ( )을 존중하는 자세가 필요하다.",
     "options": ["교류", "보급", "다양성", "융화"], "answer": 2,
     "explanation": "모양·종류 등이 여러 가지로 많은 특성은 '다양성'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 인간으로서 당연히 가지는 권리.",
     "options": ["복지", "인권", "다문화", "시민의식"], "answer": 1,
     "explanation": "'인권'은 인간으로서 당연히 가지는 권리이다."},
    {"prompt": "밑줄 친 어휘의 뜻으로 알맞은 것은? — 우리나라는 <삼권분립>의 원칙을 따른다.",
     "options": ["권력을 한곳에 모음", "선거로 대표를 뽑음", "법을 집행함", "권력을 나누어 서로 견제함"], "answer": 3,
     "explanation": "'삼권분립'은 국가 권력을 나누어 서로 견제하게 하는 제도이다."},
    {"passage_intro": "[학교 역사 동아리 유적 답사 안내]\n1. 일시: 다가오는 토요일\n2. 장소: 지역 내 유적지\n3. 안내: 문화유산 해설사가 동행하여 설명합니다.\n4. 유의사항: 유적 훼손 행위는 엄격히 금지되며, 발견 시 즉시 신고해야 합니다.\n5. 참가 신청은 담당 선생님께 접수합니다.",
     "prompt": "위 안내문에서, 이번 답사에 동행하는 사람은?",
     "options": ["문화유산 해설사", "경찰관", "소방관", "지역 주민 대표"], "answer": 0,
     "explanation": "안내문 3번에 문화유산 해설사가 동행한다고 안내되어 있다."},
    {"passage_intro": None,
     "prompt": "위 안내문에서, 유적 훼손을 발견하면 어떻게 해야 하는가?",
     "options": ["사진만 찍고 넘어간다", "직접 수리한다", "즉시 신고한다", "못 본 척한다"], "answer": 2,
     "explanation": "안내문 4번에 유적 훼손 발견 시 즉시 신고해야 한다고 안내되어 있다."},
]

EXAM3 = [
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 어떤 일에 직접 관계가 있는 사람.",
     "options": ["관계자", "대리인", "당사자", "보호자"], "answer": 2,
     "explanation": "'당사자'는 어떤 일에 직접 관계가 있는 사람을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — ( ) 확인을 위해 신분증을 제시했다.",
     "options": ["본인", "타인", "대상자", "소지자"], "answer": 0,
     "explanation": "그 사람 자신을 뜻하는 말은 '본인'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 일이 있는 바로 그날.",
     "options": ["익일", "상시", "사전", "당일"], "answer": 3,
     "explanation": "'당일'은 일이 있는 바로 그날을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 이 사건은 ( ) 경찰서로 이송되었다.",
     "options": ["인근", "관할", "인접", "현지"], "answer": 1,
     "explanation": "일정한 권한으로 관리하는 구역은 '관할'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 몸에 기운이 빠지고 정신이 멍함.",
     "options": ["안도", "초조", "허탈", "담담"], "answer": 2,
     "explanation": "'허탈'은 몸에 기운이 빠지고 정신이 멍한 상태이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 졸업식에서 ( )한 마음이 들었다.",
     "options": ["착잡", "무기력", "냉담", "절실"], "answer": 0,
     "explanation": "여러 감정이 뒤섞여 어수선한 것은 '착잡'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 매우 조심스러움.",
     "options": ["대범", "완고", "겸손", "신중"], "answer": 3,
     "explanation": "'신중'은 매우 조심스러움을 뜻한다."},
    {"prompt": "밑줄 친 어휘의 뜻으로 알맞은 것은? — 그는 <우유부단>해서 결정을 잘 못 내린다.",
     "options": ["결정이 매우 빠름", "망설이기만 하고 결단력이 없음", "고집이 아주 셈", "생각이 매우 넓음"], "answer": 1,
     "explanation": "'우유부단'은 망설이기만 하고 결단력이 없음을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 화가 났지만 감정을 ( )했다.",
     "options": ["감행", "강행", "자제", "착수"], "answer": 2,
     "explanation": "자기의 감정이나 욕망을 스스로 억제하는 것은 '자제'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 제멋대로 함부로 행함.",
     "options": ["자행", "묵인", "방관", "동참"], "answer": 0,
     "explanation": "'자행'은 제멋대로 함부로 행하는 것을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 배송비는 ( )로 청구된다.",
     "options": ["포함", "제외", "한정", "별도"], "answer": 3,
     "explanation": "원래의 것과 다른 것을 뜻하는 말은 '별도'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 본문 다음에 조건이나 예외 등을 나타내는 글.",
     "options": ["명시", "단서", "준용", "특례"], "answer": 1,
     "explanation": "'단서'는 본문 뒤에 조건이나 예외를 나타내는 글이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 쌍둥이도 성격은 서로 ( ).",
     "options": ["틀리다", "낫다", "다르다", "편하다"], "answer": 2,
     "explanation": "서로 같지 않음을 뜻할 때는 '다르다'를 쓴다. '틀리다'는 셈이나 사실이 어긋날 때 쓴다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 학생 ( ) 지켜야 할 규칙이 있다.",
     "options": ["로서", "로써", "든지", "던지"], "answer": 0,
     "explanation": "지위나 자격을 나타낼 때는 '로서'를 쓴다."},
    {"prompt": "다음 뜻풀이에 해당하는 한자성어로 알맞은 것은? — 미리 준비가 되어 있으면 걱정할 것이 없음.",
     "options": ["다다익선", "일석이조", "고진감래", "유비무환"], "answer": 3,
     "explanation": "'유비무환'은 미리 준비가 되어 있으면 걱정할 것이 없다는 뜻이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 한자성어로 가장 알맞은 것은? — 몇 번을 말해도 ( )이었다.",
     "options": ["새옹지마", "우이독경", "반신반의", "십중팔구"], "answer": 1,
     "explanation": "아무리 말해도 알아듣지 못함을 뜻하는 말은 '우이독경'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 한자성어로 알맞은 것은? — 근거 없이 널리 퍼지는 소문.",
     "options": ["온고지신", "각주구검", "유언비어", "자업자득"], "answer": 2,
     "explanation": "'유언비어'는 근거 없이 널리 퍼지는 소문을 뜻한다."},
    {"prompt": "밑줄 친 한자성어의 뜻으로 알맞은 것은? — 인생사 <새옹지마>라더니 결국 좋은 결과로 이어졌다.",
     "options": ["인생의 좋고 나쁨은 예측하기 어려움", "근거 없는 소문이 퍼짐", "고생 끝에 즐거움이 옴", "마음과 마음이 서로 통함"], "answer": 0,
     "explanation": "'새옹지마'는 인생의 좋고 나쁨은 예측하기 어렵다는 뜻이다."},
    {"passage_intro": "[학생회 임원 선거 공고]\n1. 입후보 신청 대상: 전교생 본인 또는 추천받은 학생\n2. 신청 기간: 이번 주 수요일부터 금요일까지 — 기간을 넘기면 접수되지 않습니다.\n3. 신중하게 공약을 작성하여 제출하세요.\n4. 제출한 공약은 별도 심사 없이 그대로 게시됩니다.\n5. 문의: 학생회 담당 선생님께 문의 바랍니다.",
     "prompt": "위 공고에서, 입후보 신청을 할 수 있는 사람은?",
     "options": ["3학년만 가능", "반장만 가능", "성적 상위자만 가능", "본인 또는 추천받은 학생"], "answer": 3,
     "explanation": "공고 1번에 전교생 본인 또는 추천받은 학생이 신청 대상이라고 안내되어 있다."},
    {"passage_intro": None,
     "prompt": "위 공고에서, 제출한 공약은 어떻게 처리되는가?",
     "options": ["선생님이 수정한다", "별도 심사 없이 그대로 게시된다", "학생회장이 승인해야 게시된다", "심사에서 탈락하면 폐기된다"], "answer": 1,
     "explanation": "공고 4번에 제출한 공약은 별도 심사 없이 그대로 게시된다고 안내되어 있다."},
]

EXAM4 = [
    {"prompt": "다음 뜻풀이에 해당하는 한자성어로 알맞은 것은? — 다른 사람의 하찮은 언행도 자신을 수양하는 데 도움이 됨.",
     "options": ["근묵자흑", "청출어람", "마이동풍", "타산지석"], "answer": 3,
     "explanation": "'타산지석'은 다른 사람의 하찮은 언행도 자신을 수양하는 데 도움이 된다는 뜻이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 한자성어로 가장 알맞은 것은? — 그는 ( )형 인물이라는 평을 들었다.",
     "options": ["자포자기", "대기만성", "견물생심", "과유불급"], "answer": 1,
     "explanation": "크게 될 사람은 늦게 성공한다는 뜻의 말은 '대기만성'이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 열심히 준비했다. ( ) 결과는 좋지 않았다.",
     "options": ["따라서", "게다가", "하지만", "즉"], "answer": 2,
     "explanation": "앞의 내용과 반대되는 내용을 이을 때는 '하지만'을 쓴다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 앞의 내용에 대한 구체적인 예를 드는 말.",
     "options": ["예컨대", "반면에", "결국", "요컨대"], "answer": 0,
     "explanation": "구체적인 예를 들 때 쓰는 말은 '예컨대'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 시에서 말하는 이.",
     "options": ["서술자", "심상", "정서", "화자"], "answer": 3,
     "explanation": "'화자'는 시에서 말하는 이를 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — '찬란한 슬픔'은 ( )적인 표현이다.",
     "options": ["반어", "역설", "비유", "상징"], "answer": 1,
     "explanation": "겉보기에 모순되지만 진리를 담은 표현은 '역설'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 논하는 말이나 글의 취지.",
     "options": ["전제", "근거", "논지", "요지"], "answer": 2,
     "explanation": "'논지'는 논하는 말이나 글의 취지를 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 이 글은 기존 ( )을 뒤집는 주장을 폈다.",
     "options": ["통념", "반박", "논거", "요약"], "answer": 0,
     "explanation": "일반적으로 널리 통하는 개념은 '통념'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 자료를 모아서 정리한 수치.",
     "options": ["도표", "비율", "지표", "통계"], "answer": 3,
     "explanation": "'통계'는 자료를 모아서 정리한 수치이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 여름철 판매량이 ( )했다.",
     "options": ["급감", "급증", "증가", "감소"], "answer": 1,
     "explanation": "갑자기 늚을 뜻하는 말은 '급증'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 개인을 알아볼 수 있는 정보.",
     "options": ["사생활", "지식재산권", "개인정보", "익명"], "answer": 2,
     "explanation": "'개인정보'는 개인을 알아볼 수 있는 정보를 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 타인의 사진을 ( )하면 처벌받을 수 있다.",
     "options": ["도용", "인용", "노출", "유출"], "answer": 0,
     "explanation": "남의 것을 몰래 쓰는 것은 '도용'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 지켜야 할 규칙.",
     "options": ["경고", "주의", "대응", "수칙"], "answer": 3,
     "explanation": "'수칙'은 지켜야 할 규칙을 뜻한다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 화재 시 신속히 ( )해야 한다.",
     "options": ["대기", "대피", "점검", "차단"], "answer": 1,
     "explanation": "위험을 피하여 다른 곳으로 물러나는 것은 '대피'이다."},
    {"prompt": "다음 뜻풀이에 해당하는 어휘로 알맞은 것은? — 이름을 드러내지 않아도 되는 성질.",
     "options": ["상호작용", "비대면", "익명성", "여과"], "answer": 2,
     "explanation": "'익명성'은 이름을 드러내지 않아도 되는 성질이다."},
    {"prompt": "다음 문장의 빈칸에 들어갈 말로 가장 알맞은 것은? — 화가 나도 댓글을 달 때는 ( )이 필요하다.",
     "options": ["자제력", "예절", "팔로우", "실시간"], "answer": 0,
     "explanation": "감정이나 욕구를 스스로 억제하는 힘은 '자제력'이다."},
    {"prompt": "밑줄 친 어휘의 뜻으로 알맞은 것은? — 가스 <누출>을 감지하면 즉시 신고한다.",
     "options": ["가스를 새로 채움", "가스 요금을 냄", "가스관을 설치함", "액체나 기체가 새어 나감"], "answer": 3,
     "explanation": "'누출'은 액체나 기체가 새어 나가는 것을 뜻한다."},
    {"prompt": "밑줄 친 어휘의 뜻으로 알맞은 것은? — 본인 <인증> 절차를 거쳤다.",
     "options": ["이름을 숨김", "본인임을 확인함", "비밀번호를 변경함", "계정을 삭제함"], "answer": 1,
     "explanation": "'인증'은 본인임을 확인하는 것을 뜻한다."},
    {"passage_intro": "[SNS 게시물 관리 캠페인 안내]\n1. 목적: 실시간으로 퍼지는 정보를 여과 없이 믿지 않는 습관을 기르기 위함입니다.\n2. 실천: 게시물에 댓글을 달 때는 자제력을 발휘해 예절을 지킵니다.\n3. 유의: 익명성 뒤에 숨어 타인의 사생활을 노출하는 행위는 금지됩니다.\n4. 신고: 문제 게시물을 발견하면 즉시 담당 선생님께 신고합니다.\n5. 참여: 캠페인 게시물을 팔로우하면 매주 실천 미션이 제공됩니다.",
     "prompt": "위 안내문에서, 이 캠페인이 기르고자 하는 습관은?",
     "options": ["정보를 빠르게 공유하는 습관", "익명으로 활동하는 습관", "정보를 여과 없이 믿지 않는 습관", "댓글을 많이 다는 습관"], "answer": 2,
     "explanation": "안내문 1번에 정보를 여과 없이 믿지 않는 습관을 기르는 것이 목적이라고 안내되어 있다."},
    {"passage_intro": None,
     "prompt": "위 안내문에서, 금지되는 행위로 언급된 것은?",
     "options": ["타인의 사생활을 노출하는 행위", "게시물에 댓글 달기", "캠페인 팔로우하기", "문제 게시물 신고하기"], "answer": 0,
     "explanation": "안내문 3번에 타인의 사생활을 노출하는 행위는 금지된다고 안내되어 있다."},
]


if __name__ == "__main__":
    build_exam("mid1", "2026학년도 1학기 중간고사 — 문해력 15분", "1주차 ~ 8주차 (공문서·행정 문서 읽기 ~ 디자인·영상 콘텐츠)", EXAM1, "1학기_중간고사_문해력15분")
    build_exam("final1", "2026학년도 1학기 기말고사 — 문해력 15분", "9주차 ~ 16주차 (사회 ~ 문화)", EXAM2, "1학기_기말고사_문해력15분")
    build_exam("mid2", "2026학년도 2학기 중간고사 — 문해력 15분", "17주차 ~ 24주차 (대상을 가리키는 말 ~ 한자성어(1))", EXAM3, "2학기_중간고사_문해력15분")
    build_exam("final2", "2026학년도 2학기 기말고사 — 문해력 15분", "25주차 ~ 32주차 (한자성어(2) ~ SNS·온라인 소통 언어)", EXAM4, "2학기_기말고사_문해력15분")
